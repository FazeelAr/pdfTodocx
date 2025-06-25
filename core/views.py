from django.shortcuts import render, redirect
from django.http import FileResponse
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
import os
from django.conf import settings
from PIL import Image
from io import BytesIO

def pdf_to_docx_with_images(file_obj, word_path):
    doc = Document()
    # doc.add_heading('Extracted Content from PDF', level=1)

    with fitz.open(stream=file_obj.read(), filetype='pdf') as pdf:
        for page_num, page in enumerate(pdf):
            doc.add_heading(f'Page {page_num + 1}', level=2)

            # Extract text
            text = page.get_text()
            for line in text.splitlines():
                doc.add_paragraph(line)

            # Extract and insert images
            images = page.get_images(full=True)
            for img_index, img in enumerate(images):
                xref = img[0]
                base_image = pdf.extract_image(xref)
                image_bytes = base_image["image"]
                img_ext = base_image["ext"]

                # Convert image to an in-memory file
                image_stream = BytesIO(image_bytes)
                image = Image.open(image_stream)
                image_filename = f'temp_image_{page_num}_{img_index}.{img_ext}'
                image_path = os.path.join(settings.MEDIA_ROOT, image_filename)
                image.save(image_path)

                # Insert image into docx
                doc.add_picture(image_path, width=Inches(4.5))
                os.remove(image_path)  # delete temp image after adding

    doc.save(word_path)

def convert(request):
    if request.method == 'POST':
        file = request.FILES.get('file')
        name = file.name if file else None
        name=name.split('.')[0] 
        if file:
            word_filename = f'{name}.docx'
            word_path = os.path.join(settings.MEDIA_ROOT, word_filename)

            # Process PDF and generate .docx
            pdf_to_docx_with_images(file, word_path)

            request.session['word_path'] = word_filename
            return redirect('response')

    return render(request, 'core/convert.html')

def response(request):
    word_filename = request.session.get('word_path')
    return render(request, 'core/response.html', {
        'file_url': f'{settings.MEDIA_URL}{word_filename}' if word_filename else None
    })


def index(request):
    return render(request, 'core/index.html')